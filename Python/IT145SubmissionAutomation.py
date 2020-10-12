from docx import Document
from docx.shared import Inches
import shutil
import time
import pyautogui
import os


# This function is used to grab the current window, maximize it, screenshot the window, close the window,
# create a new word document, add the screenshot from the window to the word document formatted correctly,
# and save the word document in an organized naming scheme. The pyautogui is customized for my local machine to click
# on the close button. To adapt this to a new machine a user can run pyautogui.position() to find the position
# of the close button on their machine.
def screenshot_and_save():
    CurrentWindow = pyautogui.getActiveWindow()
    CurrentWindow.maximize()
    time.sleep(1)
    pyautogui.screenshot(ScreenshotSaveFile)
    pyautogui.click(2540, 10)
    NewWordDoc = Document()
    NewWordDoc.add_picture(ScreenshotSaveFile, width=Inches(6.5), height=Inches(3.66))
    WordSavePath = NewPathName + "/" + File[:-3] + "WordScreenshot.docx"
    NewWordDoc.save(WordSavePath)


# If the program requires input the program will halt for 20 seconds before continuing.
def check_input_settings():
    if InputSetting == "1":
        time.sleep(20)
    else:
        time.sleep(1)


# This selection is used to determine the type of object to submit. Sometimes a user will need to submit an entire
# folder for a class, while other submissions will only require an individual file to be screenshot. This option
# affects which of the "if" statements gets triggered and the workflow to execute
ScreenshotSetting = input("Would you like to submit a folder or an individual file? \n"
                          "Enter 1 for Chapter Folder"
                          " or 2 for Non-Chapter Folder"
                          " or 3 for an Individual File\n")

# This setting prompts the user to select whether the program will require an input to run its course.
# Some programs designed for the class execute to process finish automatically, while others prompt for input.
# A design weakness in this program is that if it needs to loop over an entire folder and a single file requires input,
# there is no way currently to single out a file in the folder that requires input, effectively increasing the runtime
# of this program substantially if an input is required.
InputSetting = input("Does this file or folder require inputs?\n"
                     "Enter 0 for no "
                     "or 1 for yes\n")

# Rejects invalid inputs until a valid one is made.
while ScreenshotSetting != "1" and ScreenshotSetting != "2" and ScreenshotSetting != "3":
    ScreenshotSetting = input("Select option 1 or 2 or 3\n")

if ScreenshotSetting == "1":
    # The class was organized by chapters so as long as a standard naming convention was followed by me I could just
    # enter the chapter folder I would like to automatically format for submission without entering the full file path.
    ChapterNum = input("Input the chapter of the files\n")
    # File path must be altered for different users
    NewPathName = "C:/users/jpawl/Desktop/IT145ProjectsScreenshots/Chapter" + ChapterNum
    # Creates a new directory specifically for this chapters screenshots.
    os.mkdir(NewPathName)
    print(f"Directory for Chapter {ChapterNum} created")
    PythonFilePath = "C:/Users/jpawl/PycharmProjects/IT145Projects/Chapter" + ChapterNum
    # Loops through a directory and creates new paths for screenshots and documents to be saved based off of file names.
    # It also copies the python file into the screenshot folder  for convenience of access.
    for File in os.listdir(PythonFilePath):
        CurrentFile = "C:/Users/jpawl/PycharmProjects/IT145Projects/Chapter" + ChapterNum + "/" + File
        ScreenshotSaveFile = NewPathName + "/" + File[:-3] + "Screenshot.png"
        CopyFileSourceLocation = PythonFilePath + "/" + File
        CopyFileTargetLocation = NewPathName + "/"
        os.startfile(CurrentFile)
        shutil.copy(CopyFileSourceLocation, CopyFileTargetLocation)
        # If the program requires input the program will halt for 20 seconds before continuing.
        check_input_settings()
        screenshot_and_save()

elif ScreenshotSetting == "2":
    # Used to submit a folder that does not follow the standard naming conventions of the class.
    FolderPath = input("Input the full path of the Folder\n")
    # This block of code is used to obtain the file name in the directory by reversing the list of the input
    # and iterating and appending characters until the directory is reached signaled by \\.
    # The list is reversed once again to its original state.
    FolderPathList = list(FolderPath)
    FolderName = []
    for character in reversed(FolderPathList):
        if character == "\\":
            break
        if character != "\\":
            FolderName.append(character)
    FolderName.reverse()
    FolderName = "".join(FolderName)
    # File path must be altered for different users
    NewPathName = "C:/users/jpawl/Desktop/IT145ProjectsScreenshots/" + FolderName
    os.mkdir(NewPathName)
    print(f"Directory for {NewPathName} created")
    # Loops through a directory and creates new paths for screenshots and documents to be saved based off of file names.
    # It also copies the python file into the screenshot folder  for convenience of access.
    for File in os.listdir(FolderPath):
        CurrentFile = FolderPath + "/" + File
        ScreenshotSaveFile = NewPathName + "/" + File[:-3] + "Screenshot.png"
        CopyFileSourceLocation = FolderPath + "/" + File
        CopyFileTargetLocation = NewPathName + "/"
        os.startfile(CurrentFile)
        shutil.copy(CopyFileSourceLocation, CopyFileTargetLocation)
        check_input_settings()
        screenshot_and_save()

else:
    # Used to submit an individual file.
    FullPath = input("Input the full path of the file\n")
    # File path must be altered for different users
    IndividualSavePath = "C:/users/jpawl/Desktop/IT145ProjectsScreenshots/IndividualFiles"
    FileName = []
    for character in reversed(FullPath):
        if character == "\\":
            break
        if character != "\\":
            FileName.append(character)
    FileName.reverse()
    FileName = "".join(FileName[:-3])
    ScreenshotSaveFile = IndividualSavePath + "/" + FileName + "Screenshot.png"
    os.startfile(FullPath)
    shutil.copy(FullPath, IndividualSavePath)

    check_input_settings()
    # Can not use screenshot_and_save since it is not a folder
    CurrentWindow = pyautogui.getActiveWindow()
    CurrentWindow.maximize()
    time.sleep(1)
    pyautogui.screenshot(ScreenshotSaveFile)
    pyautogui.click(2540, 10)
    NewWordDoc = Document()
    NewWordDoc.add_picture(ScreenshotSaveFile, width=Inches(6.5), height=Inches(3.66))
    WordSavePath = IndividualSavePath + "/" + FileName + "WordScreenshot.docx"
    NewWordDoc.save(WordSavePath)
